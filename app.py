from io import BytesIO
import os
import asyncio
from docx import Document
import streamlit as st
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew
from crewai_tools import tool
from crewai.process import Process
import requests

def verify_gemini_api_key(api_key):
    API_VERSION = 'v1'
    api_url = f"https://generativelanguage.googleapis.com/{API_VERSION}/models?key={api_key}"
    
    try:
        response = requests.get(api_url, headers={'Content-Type': 'application/json'})
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

def verify_gpt_api_key(api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Using a simple request to the models endpoint
    response = requests.get("https://api.openai.com/v1/models", headers=headers)
    
    if response.status_code == 200:
        return True
    elif response.status_code == 401:
        return False
    else:
        print(f"Unexpected status code: {response.status_code}")
        return False

def verify_scraper_api_key(api_key):
    # ScraperAPI endpoint for a simple GET request
    url = f"http://api.scraperapi.com?api_key={api_key}&url=http://httpbin.org/ip"
    
    try:
        response = requests.get(url)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        # If we get a 401 or 403 error, it means the API key is invalid
        if e.response.status_code in [401, 403]:
            return False
        else:
            # For other HTTP errors, we'll raise an exception
            raise ValueError(f"HTTP error occurred: {e}")
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

# Function to handle RAG content generation
def generate_text(llm, query,scraper_api):
    
    @tool("Amazon Search API")
    def amazon_search_tool(query: str) -> str:
        """This is Amazon Search tool. It is used to search for 
        products on Amazon given a keyword/query to search."""
        payload = {
            'api_key': scraper_api,  # Ensure the API key is loaded from the environment
            'query': query,
            'country': 'in',
            'tld': 'in'
        }
        r = requests.get('https://api.scraperapi.com/structured/amazon/search', params=payload)
        return r.json()  # Return the response as JSON

    @tool("Amazon Product API")
    def amazon_product_tool(asins: list) -> str:
        """This is Amazon Product tool. It is used to search for 
        specific products on Amazon given a list of ASINs to search."""
        results = []
        for asin in asins:
            payload = {
                'api_key': scraper_api,
                'asin': asin,
                'country': 'in',
                'tld': 'in'
            }
            r = requests.get('https://api.scraperapi.com/structured/amazon/product', params=payload)
            results.append(r.json())
        return results  # Return a list of JSON responses

    @tool("Amazon Reviews API")
    def amazon_review_tool(asins: list) -> str:
        """This is Amazon Review tool. It is used to search reviews for 
        specific products on Amazon given a list of ASINs to search."""
        results = []
        for asin in asins:
            payload = {
                'api_key': scraper_api,
                'asin': asin,
                'country': 'in',
                'tld': 'in',
                'filter_by_star': 'all_stars'
                
            }
            r = requests.get('https://api.scraperapi.com/structured/amazon/review', params=payload)
            results.append(r.json())
        return results  # Return a list of JSON responses
    inputs = {'Product Query': query}

    searcher_agent = Agent(
        role='Amazon Searcher',
        goal='To search for products on Amazon',
        backstory="""
        You are an experienced e-commerce specialist with in-depth knowledge of Amazon's platform.
        You have spent years analyzing market trends and helping customers find the best deals and products.
        Your expertise lies in efficiently searching for products based on various criteria like price, ratings, and reviews.
        You understand the nuances of Amazon's search algorithms and can quickly filter through vast amounts of data to present the most relevant results.
        """,
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter = 5
    )

    task_searcher = Task(
    description=f'Search Amazon for the provided query and return results. The query is: {query}',
    expected_output="""
    A detailed list of Amazon products that match the search query. Each product entry should only include:
    
    - ASIN 

    At the end of your response, please provide a comma-separated list of all ASINs.
    """,
    agent=searcher_agent,
    tools=[amazon_search_tool]
)

    product_agent = Agent(
        role='Amazon Product Searcher',
        goal='To search for specific products on Amazon using ASINs',
        backstory="""
        You are an experienced e-commerce specialist with in-depth knowledge of Amazon's platform.
        You have spent years analyzing market trends and helping customers find the best deals and products.
        Your expertise lies in efficiently searching for products based on various criteria like price, ratings, and reviews.
        You understand the nuances of Amazon's search algorithms and can quickly filter through vast amounts of data to present the most relevant results.
        """,
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter = 5

    )

    task_product = Task(
    description=f'Search for detailed product information on Amazon using the ASINs extracted from the previous search results.',
    expected_output="""
    A list of detailed information for each product found in the initial search. For each product, include:
    - Full Product Name
    - Current Price
    - Total Number of Reviews
    - Average Rating
    - Product Description
    - Product Features/Specifications
    - Availability Status
    
    Present the information in a clear, organized manner, grouping details for each product together.
    """,
    context = [task_searcher],
    agent=product_agent,
    tools=[amazon_product_tool],
    async_execution=True
)

    reviews_agent = Agent(
        role='Amazon Product Reviewer',
        goal='To search for specific product reviews on Amazon using ASINs',
        backstory="""
        You are an experienced e-commerce specialist with in-depth knowledge of Amazon's platform.
        You have spent years analyzing market trends and helping customers find the best deals and products.
        Your expertise lies in efficiently searching for products based on various criteria like price, ratings, and reviews.
        You understand the nuances of Amazon's search algorithms and can quickly filter through vast amounts of data to present the most relevant results.
        """,
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter = 5

    )

    task_reviews = Task(
    description=f'Search for product reviews on Amazon using the ASINs extracted from the previous search results.',
    expected_output="""
     A list of reviews for each product found. 
     """,
    context = [task_searcher],
    agent=product_agent,
    tools=[amazon_review_tool],
    async_execution=True    
)
    formatter_agent = Agent(
    role='Output Formatter',
    goal='To organize and format product information and reviews',
    backstory="""
    You are a skilled data organizer with a keen eye for presentation. Your expertise lies in taking raw data 
    about products and their reviews, and formatting it into a clear, readable structure. You ensure that 
    information is presented consistently and logically, making it easy for users to understand and compare products.
    """,
    verbose=True,
    allow_delegation=False,
    llm=llm,
    max_iter = 5

)

    task_format = Task(
    description="""
    Format the product information and reviews for each product. Ensure that the reviews are placed below 
    the product description for each product. Organize the information in a clear, consistent manner.
    Include ALL products from the search results, do not truncate or limit the number of products.
    """,
    expected_output="""
    A well-formatted list of ALL products from the search results, where each product entry includes:
    1. Detailed Product information from task_product
    2. All reviews for the product obtained from task_reviews

    Ensure consistent formatting across all products for easy readability and comparison.
    Do not limit the number of products in the output.
    """,
    context=[task_product, task_reviews],
    agent=formatter_agent
)
    
    comparison_agent = Agent(
    role='Product Comparison Specialist',
    goal='To compare and rank Amazon products based on various criteria',
    backstory="""
    You are a highly skilled product analyst with years of experience in comparing and evaluating consumer goods. 
    Your expertise lies in analyzing product specifications, pricing, value for money, and customer reviews to provide 
    comprehensive and unbiased product comparisons. You have a keen eye for detail and can quickly identify the 
    key factors that make a product stand out in its category.
    """,
    verbose=True,
    allow_delegation=False,
    llm=llm
)

    task_comparison = Task(
    description="""
    Compare all the products from task_format. Analyze each product based on the following criteria:
    1. Product technicalities (specifications and features)
    2. Customer reviews and ratings
    Give more weightage to reviews and ratings

    After thorough analysis, determine:
    - The top 3 best overall products
    - 1 best value for money product (which may or may not be in the top 3)

    Provide a detailed explanation for each selection, highlighting the key factors that led to its ranking.
    """,
    expected_output=f"""
    
    Title - {'Best '+query}
     
    1. The total number products searched : <Number of ASINs found>

    2. Top 3 Best Overall Products 
    
    All 3 should be different. Variants of a product should not be repeated. Variants include same product in different colours and/or storage size and/or RAM:
       For each product, include:
       - Product Name\n
       - URL\n
       - Key specifications and features\n
       - Price\n
       - Average rating and number of reviews\n
       - Availability Status\n
       - Detailed explanation of why it's ranked in the top 3\n

    3. Best Value for Money Product:
        For the product, include (in seperate lines):
       - Product Name\n
       - URL\n
       - Key specifications and features\n
       - Price\n
       - Average rating and number of reviews\n
       - Availability Status\n
       - Detailed explanation of why it offers the best value for money\n
       
    Ensure that each selection is justified with clear reasoning, comparing it to other products in the list.
    
    """,
    context=[task_format],
    agent=comparison_agent
)

    crew = Crew(
    agents=[searcher_agent, product_agent, reviews_agent,formatter_agent,comparison_agent],
    tasks=[task_searcher, task_product, task_reviews,task_format,task_comparison],
    verbose=2
    
)

    result = crew.kickoff(inputs=inputs)
    return result

def main():
    
    st.header('AI Amazon Product Searcher')
    validity_model= False
    validity_scraper = False    
    
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = ''
    
    with st.sidebar:
        with st.form('Gemini/OpenAI'):
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI'))
            api_key = st.text_input(f'Enter your API key', type="password")
            scraper_key = st.text_input(f'Enter your Scraper API key', type="password")
            submitted = st.form_submit_button("Submit")

        if api_key and scraper_key:
            if model == "Gemini":
                validity_model = verify_gemini_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
                    
            elif model == "OpenAI":
                validity_model = verify_gpt_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            
            validity_scraper = verify_scraper_api_key(scraper_key)
            if validity_scraper ==True:
                st.write(f"Valid ScraperAPI API key")
            else:
                st.write(f"Invalid ScraperAPI API key")
            
            
    if validity_model and validity_scraper:                
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(model='gpt-4-turbo',temperature=0.6, max_tokens=4000,api_key=api_key)
                print("OpenAI Configured")
                return llm

            llm = asyncio.run(setup_OpenAI())
            mod = 'OpemAI'

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                print("Gemini Configured")
                return llm

            llm = asyncio.run(setup_gemini())
            mod = 'Gemini'

        query = st.text_input("Enter product query:")
        st.session_state.query = query

        if st.button("Find the best product"):
            with st.spinner("Finding product..."):
                st.session_state.generated_content = generate_text(llm, query,scraper_key)

        if st.session_state.generated_content:
            st.markdown(st.session_state.generated_content)

      # Download button logic moved outside the condition
    if st.session_state.generated_content:
        doc = Document()
        doc.add_heading(st.session_state.query, 0)
        doc.add_paragraph(st.session_state.generated_content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Download as Word Document",
            data=buffer,
            file_name=f"{st.session_state.query}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
